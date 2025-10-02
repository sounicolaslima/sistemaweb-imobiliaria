import streamlit as st
from docx import Document
import io
import os
from datetime import datetime
import tempfile
import base64

# Configuração da página
st.set_page_config(page_title="Recibo de Aluguel", page_icon="🏠")
st.title("📄 Recibo de Aluguel - Villares Imóveis")

# Inicializar lista de recibos
if 'receipts' not in st.session_state:
    st.session_state.receipts = []
if 'docx_preview' not in st.session_state:
    st.session_state.docx_preview = None

# Formulário simples
with st.form("receipt_form"):
    col1, col2 = st.columns(2)
    
    with col1:
        nomeLocatario = st.text_input("Locatário(a):")
        enderecoImovel = st.text_input("Endereço do Imóvel:")
        inicioPeriodo = st.text_input("Período de:")
        finalPeriodo = st.text_input("Até:")
        vencimento = st.text_input("Vencimento:")
        limitePagamento = st.text_input("Limite Pagamento:")
    
    with col2:
        valorAluguel = st.text_input("Valor Aluguel:", value="0,00")
        valorAgua = st.text_input("Valor Água:", value="0,00")
        valorLuz = st.text_input("Valor Luz:", value="0,00")
        valorIPTU = st.text_input("Valor IPTU:", value="0,00")
        valorCondominio = st.text_input("Valor Condomínio:", value="0,00")
        valorMulta = st.text_input("Valor Multa:", value="0,00")
        desconto = st.text_input("Desconto:", value="0,00")
        data = st.text_input("Data:", value=datetime.now().strftime('%d/%m/%Y'))
    
    # Botões
    col1, col2, col3 = st.columns(3)
    with col1: gerar = st.form_submit_button("Gerar Recibo")
    with col2: duplicar = st.form_submit_button("Incluir com Mesmos Dados")
    with col3: branco = st.form_submit_button("Incluir em Branco")

# Processar botões
if gerar:
    recibo = {
        'nomeLocatario': nomeLocatario or '',
        'enderecoImovel': enderecoImovel or '',
        'inicioPeriodo': inicioPeriodo or '',
        'finalPeriodo': finalPeriodo or '',
        'vencimento': vencimento or '',
        'limitePagamento': limitePagamento or '',
        'valorAluguel': valorAluguel or '0,00',
        'valorAgua': valorAgua or '0,00',
        'valorLuz': valorLuz or '0,00',
        'valorIPTU': valorIPTU or '0,00',
        'valorCondominio': valorCondominio or '0,00',
        'valorMulta': valorMulta or '0,00',
        'desconto': desconto or '0,00',
        'data': data or ''
    }
    st.session_state.receipts.append(recibo)
    st.session_state.docx_preview = None  # Resetar preview
    st.success("Recibo adicionado!")

if duplicar and st.session_state.receipts:
    st.session_state.receipts.append(st.session_state.receipts[-1].copy())
    st.session_state.docx_preview = None  # Resetar preview
    st.success("Recibo duplicado!")

if branco:
    st.session_state.receipts.append({k: '' for k in ['nomeLocatario', 'enderecoImovel', 'inicioPeriodo', 'finalPeriodo', 'vencimento', 'limitePagamento', 'valorAluguel', 'valorAgua', 'valorLuz', 'valorIPTU', 'valorCondominio', 'valorMulta', 'desconto', 'data']})
    st.session_state.docx_preview = None  # Resetar preview
    st.success("Recibo em branco adicionado!")

# Mostrar recibos para edição
if st.session_state.receipts:
    st.subheader(f"Recibos ({len(st.session_state.receipts)})")
    
    for i, recibo in enumerate(st.session_state.receipts):
        with st.expander(f"Recibo {i+1} - {recibo.get('nomeLocatario', 'Sem nome')}"):
            col1, col2 = st.columns(2)
            with col1:
                for campo in ['nomeLocatario', 'enderecoImovel', 'inicioPeriodo', 'finalPeriodo', 'vencimento', 'limitePagamento']:
                    recibo[campo] = st.text_input(f"{campo} {i+1}", value=recibo.get(campo, ''), key=f"{campo}_{i}")
            with col2:
                for campo in ['valorAluguel', 'valorAgua', 'valorLuz', 'valorIPTU', 'valorCondominio', 'valorMulta', 'desconto', 'data']:
                    recibo[campo] = st.text_input(f"{campo} {i+1}", value=recibo.get(campo, ''), key=f"{campo}_{i}")
            
            if st.button(f"Remover {i+1}", key=f"rm_{i}"):
                st.session_state.receipts.pop(i)
                st.session_state.docx_preview = None  # Resetar preview
                st.rerun()

# Função para substituir texto no DOCX
def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

# Gerar DOCX usando o template
def gerar_docx(recibos):
    # Verificar se o template existe
    if not os.path.exists("recibo.docx"):
        st.error("❌ Arquivo recibo.docx não encontrado! Coloque o arquivo na mesma pasta do script.")
        return None
    
    # Carregar template
    doc = Document("recibo.docx")
    
    # Mapeamento de variáveis
    variaveis = {
        '{{nomeLocatario}}': 'nomeLocatario',
        '{{enderecoImovel}}': 'enderecoImovel', 
        '{{inicioPeríodo}}': 'inicioPeriodo',
        '{{finalPeríodo}}': 'finalPeriodo',
        '{{vencimento}}': 'vencimento',
        '{{limitePagamento}}': 'limitePagamento',
        '{{valorAluguel}}': 'valorAluguel',
        '{{valorAgua}}': 'valorAgua',
        '{{valorLuz}}': 'valorLuz',
        '{{valorIPTU}}': 'valorIPTU',
        '{{valorCondominio}}': 'valorCondominio',
        '{{valorMulta}}': 'valorMulta',
        '{{desconto}}': 'desconto',
        '{{data}}': 'data'
    }
    
    # Para cada recibo, substituir as variáveis
    for i, recibo in enumerate(recibos):
        # Substituir todas as variáveis
        for var_template, var_dados in variaveis.items():
            replace_text(doc, var_template, recibo.get(var_dados, ''))
        
        # Calcular e substituir total
        try:
            total = (float(recibo['valorAluguel'].replace(',', '.')) +
                    float(recibo['valorAgua'].replace(',', '.')) +
                    float(recibo['valorLuz'].replace(',', '.')) +
                    float(recibo['valorIPTU'].replace(',', '.')) +
                    float(recibo['valorCondominio'].replace(',', '.')) +
                    float(recibo['valorMulta'].replace(',', '.')) -
                    float(recibo['desconto'].replace(',', '.')))
            replace_text(doc, '{{valorTotal}}', f"{total:,.2f}".replace('.', 'X').replace(',', '.').replace('X', ','))
        except:
            replace_text(doc, '{{valorTotal}}', '0,00')
    
    # Salvar em bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Função para criar link de download
def get_download_link(file_data, filename, text):
    b64 = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-align: center; text-decoration: none; display: inline-block; border-radius: 5px;">{text}</a>'
    return href

# Botão final
if st.session_state.receipts:
    st.markdown("---")
    
    # Botão para gerar pré-visualização
    col_prev, col_limpar = st.columns([2, 1])
    
    with col_prev:
        if st.button("👁️ Visualizar DOCX", type="primary", use_container_width=True):
            with st.spinner("🔄 Gerando pré-visualização do DOCX..."):
                doc_io = gerar_docx(st.session_state.receipts)
                
                if doc_io:
                    st.session_state.docx_preview = doc_io.getvalue()
                    st.balloons()
                    st.success("✅ DOCX gerado com sucesso!")
    
    with col_limpar:
        if st.button("🗑️ Limpar Tudo", use_container_width=True):
            st.session_state.receipts = []
            st.session_state.docx_preview = None
            st.rerun()
    
    # Mostrar pré-visualização se existir
    if st.session_state.docx_preview:
        st.markdown("---")
        st.subheader("📋 Pré-visualização do DOCX")
        
        # Informações do documento
        st.info(f"📄 **Documento gerado com {len(st.session_state.receipts)} recibo(s)**")
        
        # Botão de download
        st.markdown(
            get_download_link(
                st.session_state.docx_preview,
                f"recibos_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                "⬇️ Baixar DOCX"
            ),
            unsafe_allow_html=True
        )
        
        # Mostrar miniaturas dos recibos (visualização simplificada)
        st.subheader("👀 Visualização dos Recibos no Documento:")
        
        for i, recibo in enumerate(st.session_state.receipts):
            with st.expander(f"📋 Recibo {i+1} no DOCX", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Dados do Recibo:**")
                    st.write(f"**Locatário:** {recibo.get('nomeLocatario', 'Não informado')}")
                    st.write(f"**Endereço:** {recibo.get('enderecoImovel', 'Não informado')}")
                    st.write(f"**Período:** {recibo.get('inicioPeriodo', 'Não informado')} a {recibo.get('finalPeriodo', 'Não informado')}")
                    st.write(f"**Vencimento:** {recibo.get('vencimento', 'Não informado')}")
                
                with col2:
                    st.write("**Valores:**")
                    st.write(f"**Aluguel:** R$ {recibo.get('valorAluguel', '0,00')}")
                    st.write(f"**Água:** R$ {recibo.get('valorAgua', '0,00')}")
                    st.write(f"**Luz:** R$ {recibo.get('valorLuz', '0,00')}")
                    st.write(f"**IPTU:** R$ {recibo.get('valorIPTU', '0,00')}")
                    st.write(f"**Condomínio:** R$ {recibo.get('valorCondominio', '0,00')}")
                    st.write(f"**Data:** {recibo.get('data', 'Não informado')}")
                
                # Calcular total
                try:
                    total = (float(recibo['valorAluguel'].replace(',', '.')) +
                            float(recibo['valorAgua'].replace(',', '.')) +
                            float(recibo['valorLuz'].replace(',', '.')) +
                            float(recibo['valorIPTU'].replace(',', '.')) +
                            float(recibo['valorCondominio'].replace(',', '.')) +
                            float(recibo['valorMulta'].replace(',', '.')) -
                            float(recibo['desconto'].replace(',', '.')))
                    st.success(f"**TOTAL: R$ {total:,.2f}**".replace('.', 'X').replace(',', '.').replace('X', ','))
                except:
                    st.warning("**TOTAL:** Não foi possível calcular")
        
        st.info("💡 **Dica:** O documento DOCX gerado contém todos os recibos no formato original do template, prontos para impressão.")

# Instruções
st.markdown("---")
st.info("""
💡 **Como usar:**
1. Coloque o arquivo **recibo.docx** na mesma pasta
2. Preencha os dados e clique em **Gerar Recibo**
3. Use os botões para adicionar mais recibos (até 3 por página)
4. Clique em **👁️ Visualizar DOCX** para gerar o documento
5. **Verifique a pré-visualização** dos dados
6. Clique em **⬇️ Baixar DOCX** para salvar o arquivo

📌 **Formato:** O DOCX gerado manterá exatamente o layout do seu template com todos os recibos!
""")

# Verificação do template
if not os.path.exists("recibo.docx"):
    st.error("""
    ❌ **Template não encontrado!**
    
    Para usar este sistema, você precisa:
    1. Colocar o arquivo **recibo.docx** na mesma pasta deste script
    2. Certificar-se que o arquivo contém as variáveis {{nomeLocatario}}, {{valorAluguel}}, etc.
    3. Reiniciar o aplicativo após colocar o arquivo
    """)
else:
    st.success("✅ Template recibo.docx encontrado e pronto para uso!")